from __future__ import annotations

import json
import re
from typing import Any
from urllib.parse import urlparse

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover - environment dependent
    Document = None

from app.utils.text_utils import lexical_overlap_score, normalize_text, tokenize


class KnowledgeBaseService:
    RESOURCE_MARKERS = [
        "بودكاست", "youtube", "youtu.be", "spotify", "كتاب", "تلخيص كتاب", "شرح كتاب",
        "بوكافيين", "دوباميكافين", "دوباماكفيين", "فنجان", "جلسة", "افيدونا", "الدحيح", "الحلقة", "حلقة",
        "pbs", "ted", "mustafa hosny", "amel atia", "gottman", "episode", "podcast",
    ]
    PLATFORM_VIDEO_MARKERS = [
        "youtube", "youtu.be", "youtube.com", "الحلقة", "حلقة", "الدحيح", "ted", "pbs",
        "بوكافيين", "دوباميكافين", "دوباماكفيين", "mustafa hosny", "امل عطية", "امال عطية",
    ]
    PODCAST_MARKERS = ["بودكاست", "spotify", "فنجان", "جلسة", "افيدونا", "podcast"]
    BOOK_TITLE_MARKERS = [
        "كتاب", "السماح بالرحيل", "ازالة القلق", "إزالة القلق", "ربما ينبغي ان تكلم احدا", "ربما ينبغي أن تكلم أحداً",
        "فن اللامبالاة", "الدردشة العقلية", "وهم المقارنه", "وهم المقارنة", "قوة الان", "قوه الان",
        "الاثنى عشر اسبوع", "الاثني عشر اسبوع", "امة الدوبامين", "أمة الدوبامين", "الدوبامين",
        "دع القلق وابدأ الحياة", "دع القلق وابدا الحياه", "النوع الاخر من الذكاء",
    ]
    KNOWN_BOOK_CHANNELS = ["بوكافيين", "دوباميكافين", "دوباماكفيين"]
    NOISE_RESOURCE_MARKERS = [
        "grammarly", "ai built for school", "raise grades", "speed up work",
    ]
    DESCRIPTIVE_CUT_MARKERS = [
        " وهي ", " وهو ", " حيث ", " موضح", " موضحة", " يركز ", " تركز ", " وينتقد ", " وتقدم ", " تقدم ",
        " الفائدة ", " وهذه ", " وهذا ", " مناسبة ل", " يرتبط ", " ترتبط ", " تناقش ",
    ]
    GENERIC_PREFIXES = [
        "حلقات عن", "من قناة", "برنامج", "سلسلة", "المصادر العربية", "المصادر الأجنبية", "المصادر الاجنبية",
    ]
    MBTI_TYPES = [
        "INFJ", "INFP", "ENFJ", "ENFP",
        "INTJ", "INTP", "ENTJ", "ENTP",
        "ISTJ", "ISFJ", "ESTJ", "ESFJ",
        "ISTP", "ISFP", "ESTP", "ESFP",
    ]

    def __init__(self, settings) -> None:
        self.settings = settings
        self.knowledge: dict[str, Any] = {"mbti": {}, "emotion": {}}
        self.chunks: list[dict[str, Any]] = []

    def load(self) -> None:
        if self.settings.KB_JSON_PATH.exists() and self.settings.CHUNKS_JSON_PATH.exists():
            self.knowledge = json.loads(self.settings.KB_JSON_PATH.read_text(encoding="utf-8"))
            self.chunks = json.loads(self.settings.CHUNKS_JSON_PATH.read_text(encoding="utf-8"))
            return
        self.rebuild_from_sources()

    def rebuild_from_sources(self) -> None:
        if Document is None:
            if self.settings.KB_JSON_PATH.exists() and self.settings.CHUNKS_JSON_PATH.exists():
                self.load()
                return
            raise RuntimeError("python-docx is required to rebuild knowledge from .docx sources.")
        mbti = self._parse_mbti_doc(self.settings.MBTI_DOCX_PATH)
        emotion = self._parse_emotion_doc(self.settings.EMOTION_DOCX_PATH)
        self.knowledge = {"mbti": mbti, "emotion": emotion}
        self.chunks = self._build_chunks(self.knowledge)
        self.settings.KB_JSON_PATH.parent.mkdir(parents=True, exist_ok=True)
        self.settings.CHUNKS_JSON_PATH.parent.mkdir(parents=True, exist_ok=True)
        self.settings.KB_JSON_PATH.write_text(json.dumps(self.knowledge, ensure_ascii=False, indent=2), encoding="utf-8")
        self.settings.CHUNKS_JSON_PATH.write_text(json.dumps(self.chunks, ensure_ascii=False, indent=2), encoding="utf-8")

    def get_mbti_overview(self, mbti_type: str | None) -> dict[str, Any]:
        if not mbti_type:
            return {"mbti_type": None, "core_problems": [], "consequences": [], "issues": []}
        return self.knowledge.get("mbti", {}).get(mbti_type, {"mbti_type": mbti_type, "core_problems": [], "consequences": [], "issues": []})

    def merge_advice(
        self,
        *,
        mbti_type: str | None,
        issue_title: str | None = None,
        topic_title: str | None = None,
        question_title: str | None = None,
    ) -> list[str]:
        merged: list[str] = []
        seen: set[str] = set()
        for source in [
            self.get_issue_advice(mbti_type, issue_title),
            self.get_emotion_advice(topic_title=topic_title, question_title=question_title),
        ]:
            for line in source:
                clean = line.strip()
                key = self._normalize_label(clean)
                if not clean or not key or key in seen:
                    continue
                seen.add(key)
                merged.append(clean)

        prioritized = [
            line for line in merged
            if self._looks_like_action_step(line) or line.startswith("قاعدة") or "خطوة" in line or "تدريب" in line
        ]
        if prioritized:
            return self._dedupe_similar(prioritized + [line for line in merged if line not in prioritized])[:6]
        return self._dedupe_similar(merged)[:6]

    def get_issue_advice(self, mbti_type: str | None, issue_title: str | None = None) -> list[str]:
        if not mbti_type:
            return []
        issues = self.knowledge.get("mbti", {}).get(mbti_type, {}).get("issues", [])
        if issue_title:
            for issue in issues:
                if self._normalize_issue_title(issue.get("title", "")) == self._normalize_issue_title(issue_title):
                    return self._clean_advice_lines(issue.get("advice", []))
        for issue in issues:
            advice = self._clean_advice_lines(issue.get("advice", []))
            if advice:
                return advice
        return []

    def get_emotion_entry(self, *, topic_title: str | None = None, question_title: str | None = None) -> dict[str, Any] | None:
        topics = self.knowledge.get("emotion", {})
        normalized_topic = self._normalize_label(topic_title or "")
        normalized_question = self._normalize_label(question_title or "")
        for topic in topics.values():
            if normalized_topic and self._normalize_label(topic.get("topic_title", "")) != normalized_topic:
                continue
            for question in topic.get("questions", []):
                if normalized_question and self._normalize_label(question.get("question", "")) != normalized_question:
                    continue
                return question
        if normalized_topic:
            topic = topics.get(topic_title or "")
            if topic and topic.get("questions"):
                return topic["questions"][0]
        for topic in topics.values():
            if topic.get("questions"):
                return topic["questions"][0]
        return None

    def get_emotion_advice(self, *, topic_title: str | None = None, question_title: str | None = None) -> list[str]:
        topics = self.knowledge.get("emotion", {})
        normalized_topic = self._normalize_label(topic_title or "")
        normalized_question = self._normalize_label(question_title or "")
        for topic in topics.values():
            if normalized_topic and self._normalize_label(topic.get("topic_title", "")) != normalized_topic:
                continue
            for question in topic.get("questions", []):
                if normalized_question and self._normalize_label(question.get("question", "")) != normalized_question:
                    continue
                steps = self._clean_advice_lines(question.get("steps", []))
                if steps:
                    return steps
        for topic in topics.values():
            for question in topic.get("questions", []):
                steps = self._clean_advice_lines(question.get("steps", []))
                if steps:
                    return steps
        return []

    def recommend_resources(
        self,
        mbti_type: str | None,
        query: str,
        *,
        limit: int = 6,
        category: str | None = None,
        issue_titles: list[str] | None = None,
        topic_titles: list[str] | None = None,
        question_titles: list[str] | None = None,
        require_url: bool = False,
    ) -> list[dict[str, Any]]:
        issue_titles = issue_titles or []
        topic_titles = topic_titles or []
        question_titles = question_titles or []
        norm_issue_titles = {self._normalize_issue_title(value) for value in issue_titles if value}
        norm_topic_titles = {self._normalize_label(value) for value in topic_titles if value}
        norm_question_titles = {self._normalize_label(value) for value in question_titles if value}

        query_tokens = set(tokenize(query))
        candidates = self._candidate_resources(
            mbti_type=mbti_type,
            issue_titles=issue_titles,
            topic_titles=topic_titles,
            question_titles=question_titles,
        )
        if not candidates:
            candidates = self._all_resources()

        scored: list[tuple[float, dict[str, Any]]] = []
        for item in candidates:
            if category and item.get("category") != category:
                continue
            if require_url and not self._has_valid_url(item.get("url")):
                continue
            if item.get("source_collection") == "mbti" and mbti_type and item.get("mbti_type") not in {None, mbti_type}:
                continue
            if self._looks_like_noise_resource(item.get("title", "")):
                continue

            title = (item.get("title") or "").strip()
            if not title:
                continue

            issue_title = item.get("issue_title") or ""
            topic_title = item.get("topic_title") or ""
            query_overlap = lexical_overlap_score(query, title)
            issue_overlap = lexical_overlap_score(query, issue_title)
            topic_overlap = lexical_overlap_score(query, topic_title)

            score = 0.0
            if self._normalize_issue_title(issue_title) in norm_issue_titles and issue_title:
                score += 3.0
            if self._normalize_label(topic_title) in norm_topic_titles and topic_title:
                score += 2.4
            if self._normalize_label(issue_title) in norm_question_titles and issue_title:
                score += 2.2
            if item.get("source_collection") == "mbti" and mbti_type and item.get("mbti_type") == mbti_type:
                score += 0.35
            if self._has_valid_url(item.get("url")):
                score += 0.35
            score += query_overlap * 2.4
            score += issue_overlap * 1.4
            score += topic_overlap * 1.0

            title_tokens = set(tokenize(title))
            issue_tokens = set(tokenize(issue_title))
            topic_tokens = set(tokenize(topic_title))
            shared = len(query_tokens & (title_tokens | issue_tokens | topic_tokens))
            score += min(shared * 0.18, 0.9)

            if category == "book" and not self._is_true_book(item):
                continue
            if category in {"video", "podcast"} and not self._has_valid_url(item.get("url")):
                continue
            if score < 0.7:
                continue

            scored.append((score, item))

        scored.sort(key=lambda row: row[0], reverse=True)

        selected: list[dict[str, Any]] = []
        seen: set[str] = set()
        for score, item in scored:
            key = self._resource_key(item)
            if key in seen:
                continue
            seen.add(key)
            selected.append({
                **item,
                "score": round(score, 4),
                "why_recommended": self._build_resource_reason(
                    item,
                    query_tokens=query_tokens,
                    norm_issue_titles=norm_issue_titles,
                    norm_topic_titles=norm_topic_titles,
                    norm_question_titles=norm_question_titles,
                ),
            })
            if len(selected) >= limit:
                break
        return selected

    def _candidate_resources(
        self,
        *,
        mbti_type: str | None,
        issue_titles: list[str],
        topic_titles: list[str],
        question_titles: list[str],
    ) -> list[dict[str, Any]]:
        norm_issue_titles = {self._normalize_issue_title(value) for value in issue_titles if value}
        norm_topic_titles = {self._normalize_label(value) for value in topic_titles if value}
        norm_question_titles = {self._normalize_label(value) for value in question_titles if value}

        candidates: list[dict[str, Any]] = []
        seen: set[str] = set()

        def add_many(items: list[dict[str, Any]]) -> None:
            for item in items:
                key = self._resource_key(item)
                if key and key not in seen:
                    seen.add(key)
                    candidates.append(item)

        if mbti_type:
            mbti_entry = self.knowledge.get("mbti", {}).get(mbti_type, {})
            add_many(mbti_entry.get("resources", []))
            for issue in mbti_entry.get("issues", []):
                if not norm_issue_titles or self._normalize_issue_title(issue.get("title", "")) in norm_issue_titles:
                    add_many(issue.get("resources", []))

        for topic in self.knowledge.get("emotion", {}).values():
            topic_match = self._normalize_label(topic.get("topic_title", "")) in norm_topic_titles if norm_topic_titles else False
            if topic_match:
                add_many(topic.get("resources", []))
            for question in topic.get("questions", []):
                question_match = self._normalize_label(question.get("question", "")) in norm_question_titles if norm_question_titles else False
                inherited_issue_match = self._normalize_issue_title(question.get("question", "")) in norm_issue_titles if norm_issue_titles else False
                if topic_match or question_match or inherited_issue_match:
                    add_many(question.get("resources", []))

        return candidates

    def _all_resources(self) -> list[dict[str, Any]]:
        resources: list[dict[str, Any]] = []
        for mbti_data in self.knowledge.get("mbti", {}).values():
            resources.extend(mbti_data.get("resources", []))
            for issue in mbti_data.get("issues", []):
                resources.extend(issue.get("resources", []))
        for topic_data in self.knowledge.get("emotion", {}).values():
            resources.extend(topic_data.get("resources", []))
            for question in topic_data.get("questions", []):
                resources.extend(question.get("resources", []))
        return resources

    def _parse_mbti_doc(self, path) -> dict[str, Any]:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        indices = [i for i, line in enumerate(paragraphs) if self._is_mbti_heading(line)]
        knowledge: dict[str, Any] = {}
        for idx, start in enumerate(indices):
            end = indices[idx + 1] if idx + 1 < len(indices) else len(paragraphs)
            mbti = self._normalize_type(paragraphs[start])
            lines = [line for line in paragraphs[start + 1:end] if line.strip()]
            parsed = self._parse_mbti_block(lines, mbti)
            parsed["resources"] = self._postprocess_resources(
                parsed["resources"],
                source_collection="mbti",
                source_document=self.settings.MBTI_DOCX_PATH.name,
                mbti_type=mbti,
            )
            for issue in parsed["issues"]:
                issue["resources"] = self._postprocess_resources(
                    issue.get("resources", []),
                    issue_title=issue["title"],
                    source_collection="mbti",
                    source_document=self.settings.MBTI_DOCX_PATH.name,
                    mbti_type=mbti,
                    topic_title=mbti,
                )
            knowledge[mbti] = parsed
        return knowledge

    def _parse_mbti_block(self, lines: list[str], mbti: str) -> dict[str, Any]:
        def idx_of(predicate):
            for i, value in enumerate(lines):
                if predicate(value):
                    return i
            return -1

        problems_idx = idx_of(lambda value: "المشكلات الجوهرية" in value)
        results_idx = idx_of(lambda value: value.startswith("النتائج"))
        solutions_idx = idx_of(lambda value: "الحلول" in value)

        problems: list[str] = []
        consequences: list[str] = []
        solution_lines: list[str] = []

        if problems_idx != -1 and results_idx != -1:
            problems = [line for line in lines[problems_idx + 1:results_idx] if line]
            if solutions_idx != -1:
                consequences = [line for line in lines[results_idx + 1:solutions_idx] if line]
                solution_lines = [line for line in lines[solutions_idx + 1:] if line]
            else:
                consequences = [line for line in lines[results_idx + 1:results_idx + 5] if line]
                solution_lines = [line for line in lines[results_idx + 5:] if line]
        else:
            solution_lines = [line for line in lines if line]

        issues: list[dict[str, Any]] = []
        global_resources: list[dict[str, Any]] = []
        current_issue: dict[str, Any] | None = None

        def flush_issue() -> None:
            nonlocal current_issue
            if current_issue:
                current_issue["advice"] = self._clean_advice_lines(current_issue.get("advice", []))
                issues.append(current_issue)
                current_issue = None

        for raw in solution_lines:
            line = raw.strip()
            normalized_issue = self._normalize_issue_title(line)
            if line.startswith("المشكلة:") or ("المشكلة:" in line and len(normalized_issue) > 3):
                flush_issue()
                current_issue = {"title": normalized_issue, "advice": [], "resources": []}
                continue
            if current_issue is None:
                if self._looks_like_resource(line):
                    global_resources.append({"title": line})
                continue
            if self._looks_like_resource(line):
                current_issue["resources"].append({"title": line, "issue_title": current_issue["title"]})
            else:
                current_issue["advice"].append(line)
        flush_issue()

        return {
            "mbti_type": mbti,
            "heading": mbti,
            "core_problems": [self._normalize_issue_title(item) for item in problems if item],
            "consequences": [item.strip() for item in consequences if item],
            "issues": issues,
            "resources": global_resources,
        }

    def _parse_emotion_doc(self, path) -> dict[str, Any]:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        topics: dict[str, Any] = {}
        current_topic: dict[str, Any] | None = None
        current_question: dict[str, Any] | None = None

        def ensure_topic(title: str) -> dict[str, Any]:
            if title not in topics:
                topics[title] = {
                    "topic_title": title,
                    "questions": [],
                    "resources": [],
                    "document": self.settings.EMOTION_DOCX_PATH.name,
                }
            return topics[title]

        def flush_question() -> None:
            nonlocal current_question, current_topic
            if current_question and current_topic:
                current_question["steps"] = self._clean_advice_lines(current_question.get("steps", []))
                current_topic["questions"].append(current_question)
            current_question = None

        for line in paragraphs:
            if self._is_emotion_topic_heading(line):
                flush_question()
                topic_title = self._clean_topic_title(line)
                current_topic = ensure_topic(topic_title)
                continue

            if self._is_question_line(line):
                flush_question()
                if current_topic is None:
                    current_topic = ensure_topic("عام")
                question, answer_intro = self._split_question_block(line)
                current_question = {
                    "question": question,
                    "answer_intro": answer_intro,
                    "details": [],
                    "steps": [],
                    "resources": [],
                    "topic_title": current_topic["topic_title"],
                }
                continue

            if self._looks_like_resource(line):
                target = current_question["resources"] if current_question else (current_topic["resources"] if current_topic else None)
                if target is not None:
                    target.append({
                        "title": line,
                        "topic_title": current_topic["topic_title"] if current_topic else None,
                        "issue_title": current_question["question"] if current_question else None,
                    })
                continue

            if current_question is not None:
                current_question["details"].append(line)
                if self._looks_like_action_step(line):
                    current_question["steps"].append(line)

        flush_question()

        for topic in topics.values():
            topic["resources"] = self._postprocess_resources(
                topic.get("resources", []),
                source_collection="emotion",
                source_document=self.settings.EMOTION_DOCX_PATH.name,
                topic_title=topic["topic_title"],
            )
            for question in topic.get("questions", []):
                question["resources"] = self._postprocess_resources(
                    question.get("resources", []),
                    issue_title=question["question"],
                    source_collection="emotion",
                    source_document=self.settings.EMOTION_DOCX_PATH.name,
                    topic_title=topic["topic_title"],
                )
        return topics

    def _build_chunks(self, knowledge: dict[str, Any]) -> list[dict[str, Any]]:
        chunks: list[dict[str, Any]] = []
        for mbti_type, item in knowledge.get("mbti", {}).items():
            chunks.append({
                "id": f"{mbti_type}-overview",
                "domain": "mbti",
                "chunk_type": "overview",
                "title": mbti_type,
                "topic_title": None,
                "mbti_type": mbti_type,
                "text": (
                    f"MBTI: {mbti_type}\n"
                    f"المشكلات الجوهرية: {' | '.join(item.get('core_problems', []))}\n"
                    f"النتائج: {' | '.join(item.get('consequences', []))}"
                ),
            })
            for idx, issue in enumerate(item.get("issues", []), start=1):
                resource_titles = [res.get("title", "") for res in issue.get("resources", [])[:5]]
                chunks.append({
                    "id": f"{mbti_type}-issue-{idx}",
                    "domain": "mbti",
                    "chunk_type": "issue",
                    "title": issue.get("title", f"Issue {idx}"),
                    "topic_title": mbti_type,
                    "mbti_type": mbti_type,
                    "text": (
                        f"MBTI: {mbti_type}\n"
                        f"Issue: {issue.get('title', '')}\n"
                        f"Advice: {' | '.join(issue.get('advice', [])[:6])}\n"
                        f"Resources: {' | '.join(resource_titles)}"
                    ),
                })
        for topic_title, topic in knowledge.get("emotion", {}).items():
            sample_questions = [q.get("question", "") for q in topic.get("questions", [])[:6]]
            chunks.append({
                "id": f"emotion-topic-{self._slugify(topic_title)}",
                "domain": "emotion",
                "chunk_type": "emotion_topic",
                "title": topic_title,
                "topic_title": topic_title,
                "mbti_type": None,
                "text": f"Topic: {topic_title}\nQuestions: {' | '.join(sample_questions)}",
            })
            for idx, question in enumerate(topic.get("questions", []), start=1):
                resource_titles = [res.get("title", "") for res in question.get("resources", [])[:4]]
                chunks.append({
                    "id": f"emotion-question-{self._slugify(topic_title)}-{idx}",
                    "domain": "emotion",
                    "chunk_type": "emotion_question",
                    "title": question.get("question", f"Question {idx}"),
                    "topic_title": topic_title,
                    "mbti_type": None,
                    "text": (
                        f"Topic: {topic_title}\n"
                        f"Question: {question.get('question', '')}\n"
                        f"Answer: {question.get('answer_intro', '')}\n"
                        f"Details: {' | '.join(question.get('details', [])[:5])}\n"
                        f"Steps: {' | '.join(question.get('steps', [])[:5])}\n"
                        f"Resources: {' | '.join(resource_titles)}"
                    ),
                })
        return chunks

    def _is_mbti_heading(self, text: str) -> bool:
        return bool(re.match(r"^(%s)\b" % "|".join(self.MBTI_TYPES), text.strip()))

    def _normalize_type(self, heading: str) -> str:
        match = re.match(r"^(%s)\b" % "|".join(self.MBTI_TYPES), heading.strip())
        if not match:
            raise ValueError(f"Invalid MBTI heading: {heading}")
        return match.group(1)

    def _extract_url(self, text: str) -> str | None:
        match = re.search(r"https?://\S+", text)
        return match.group(0).rstrip(").,؛،]") if match else None

    def _remove_url_from_text(self, text: str) -> str:
        return re.sub(r"https?://\S+", "", text).strip(" -—–:؛،\n\t")

    def _contains_word(self, text: str, marker: str) -> bool:
        text_lower = text.lower()
        marker_lower = marker.lower()
        if " " in marker_lower:
            return marker_lower in text_lower
        return bool(re.search(rf"(?<!\w){re.escape(marker_lower)}(?!\w)", text_lower))

    def _classify_resource(self, text: str, url: str | None = None) -> str:
        lowered = text.lower()
        url_lower = (url or "").lower()
        if any(noise in lowered for noise in self.NOISE_RESOURCE_MARKERS):
            return "general"
        if "spotify" in url_lower or any(self._contains_word(text, marker) for marker in self.PODCAST_MARKERS):
            return "podcast"
        if any(marker in url_lower for marker in ["youtube.com", "youtu.be"]):
            return "video"
        if any(self._contains_word(text, marker) for marker in self.PLATFORM_VIDEO_MARKERS):
            return "video"
        if "شرح كتاب" in text or "تلخيص كتاب" in text:
            return "video"
        if self._is_true_book({"title": text, "url": url}):
            return "book"
        return "general"

    def _is_true_book(self, item: dict[str, Any]) -> bool:
        text = (item.get("title") or "").strip().lower()
        url = (item.get("url") or "").lower()
        if not text:
            return False
        if self._looks_like_action_step(text) or self._is_question_line(text):
            return False
        if any(marker in url for marker in ["youtube.com", "youtu.be", "spotify"]):
            return False
        if any(channel.lower() in text for channel in self.KNOWN_BOOK_CHANNELS):
            return False
        if any(marker in text for marker in ["youtube", "youtu", "spotify", "بودكاست", "الحلقة", "حلقة", "episode", "pbs", "ted"]):
            return False
        normalized = self._normalize_label(text)
        if "كتاب" in normalized and len(normalized.split()) <= 10:
            return True
        return any(self._normalize_label(marker) in normalized for marker in self.BOOK_TITLE_MARKERS)

    def _looks_like_noise_resource(self, text: str) -> bool:
        normalized = self._normalize_label(text)
        return any(self._normalize_label(marker) in normalized for marker in self.NOISE_RESOURCE_MARKERS)

    def _looks_like_resource(self, text: str) -> bool:
        lowered = text.lower().strip()
        if lowered.startswith("http"):
            return True
        if self._looks_like_noise_resource(text):
            return True
        if self._looks_like_action_step(text) or self._is_question_line(text):
            return False
        if any(self._contains_word(text, marker) for marker in self.RESOURCE_MARKERS):
            return True
        if re.search(r"https?://|youtu\.be|youtube|spotify", lowered):
            return True
        if ("|" in text or "#" in text) and len(text.split()) <= 30:
            return True
        return False

    def _looks_like_action_step(self, text: str) -> bool:
        clean = text.strip()
        lowered = clean.lower()
        starters = [
            "خدي", "حددي", "اركزي", "لاحظي", "اسألي", "اكتبي", "اتواصلي", "اطلبي", "حاولي",
            "ابدئي", "امشي", "شاركي", "قولي", "راجعي", "سيبيه", "تذكري", "قللي", "قسمي", "قسّمي",
            "تجنبي", "اتفقي", "افتكري", "اختاري", "خلي", "أخدي", "ابعدي", "اعملي", "خليكي", "تعملي",
        ]
        if any(clean.startswith(word) for word in starters):
            return True
        if lowered.startswith("الافضل") or lowered.startswith("الأفضل"):
            return True
        return False

    def _is_emotion_topic_heading(self, text: str) -> bool:
        stripped = text.strip()
        return stripped.startswith("(") and stripped.endswith(")") and len(stripped) <= 120

    def _clean_topic_title(self, text: str) -> str:
        return text.strip().strip("() ")

    def _is_question_line(self, text: str) -> bool:
        clean = text.strip()
        if "؟" not in clean:
            return False
        if len(clean.split()) > 40:
            return False
        starters = ["أنا", "ليه", "إزاي", "ازاي", "هل", "هو", "إمتى", "امتي"]
        return any(clean.startswith(starter) for starter in starters)

    def _split_question_block(self, text: str) -> tuple[str, str]:
        lines = [part.strip() for part in text.splitlines() if part.strip()]
        question = lines[0]
        answer_intro = " ".join(lines[1:]).strip()
        return question, answer_intro

    def _normalize_issue_title(self, text: str) -> str:
        cleaned = text.replace("المشكلة:", "").replace("✔", " ")
        cleaned = re.sub(r"^\d+\\\s*", "", cleaned)
        cleaned = re.sub(r"(?<!\w)الحل(?=!|:|\s|$)", " ", cleaned)
        cleaned = re.sub(r"(?<!\w)الحلول(?=!|:|\s|$)", " ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" :\n\t")
        return cleaned

    def _normalize_label(self, text: str) -> str:
        text = text or ""
        text = text.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا")
        text = text.replace("ة", "ه").replace("ى", "ي")
        text = re.sub(r"[^\w\s\u0600-\u06FF]+", " ", text.lower())
        return re.sub(r"\s+", " ", text).strip()

    def _clean_advice_lines(self, lines: list[str]) -> list[str]:
        cleaned: list[str] = []
        seen: set[str] = set()
        for line in lines:
            item = re.sub(r"\s+", " ", line).strip("-•* ")
            if not item:
                continue
            lowered = item.lower()
            if lowered in {"الحل:", "الحلول:", "✔ الحل:", "الحل", "الحلول", "الإجابة:"}:
                continue
            resource_like_markers = [
                "بودكاست", "كتاب", "تلخيص كتاب", "شرح كتاب", "youtube", "youtu", "spotify",
                "بوكافيين", "دوباميكافين", "دوباماكفيين", "فيديو", "الحلقة", "حلقة", "pbs", "ted"
            ]
            if self._looks_like_resource(item) or any(marker.lower() in lowered for marker in resource_like_markers):
                continue
            if len(item) < 6 or len(item.split()) > 18:
                continue
            if len(item.split()) <= 5 and not self._looks_like_action_step(item) and not item.startswith("قاعدة"):
                continue
            key = self._normalize_label(item)
            if key not in seen:
                seen.add(key)
                cleaned.append(item)
        return self._dedupe_similar(cleaned)

    def _resource_key(self, item: dict[str, Any]) -> str:
        return (
            f"{(item.get('title') or '').strip().lower()}|{(item.get('url') or '').strip().lower()}|"
            f"{item.get('category') or ''}|{item.get('source_collection') or ''}"
        )

    def _build_resource_reason(
        self,
        item: dict[str, Any],
        *,
        query_tokens: set[str],
        norm_issue_titles: set[str],
        norm_topic_titles: set[str],
        norm_question_titles: set[str],
    ) -> str:
        issue_title = item.get("issue_title") or ""
        topic_title = item.get("topic_title") or ""
        title_tokens = set(tokenize(item.get("title", "")))
        query_hits = sorted(token for token in (query_tokens & title_tokens) if len(token) >= 3)[:3]

        reasons: list[str] = []
        if self._normalize_issue_title(issue_title) in norm_issue_titles and issue_title:
            reasons.append(f"مرتبط مباشرة بمحور: {self._normalize_issue_title(issue_title)}")
        elif self._normalize_label(issue_title) in norm_question_titles and issue_title:
            reasons.append(f"مرتبط مباشرة بالسؤال الأقرب: {issue_title}")
        elif self._normalize_label(topic_title) in norm_topic_titles and topic_title:
            reasons.append(f"مرتبط مباشرة بموضوع: {topic_title}")

        if query_hits:
            reasons.append(f"وفي العنوان كلمات قريبة من سؤالك: {', '.join(query_hits)}")
        elif item.get("category") == "video":
            reasons.append("مناسب كخطوة سريعة مع رابط مباشر")
        elif item.get("category") == "podcast":
            reasons.append("مناسب لو حابة شرحًا صوتيًا أهدأ مع رابط مباشر")
        elif item.get("category") == "book":
            reasons.append("مناسب لو الهدف فهم أعمق من مصدر كتاب")

        return "، ".join(reasons) if reasons else "ظهر لأنه الأقرب ارتباطًا بالسؤال داخل البيانات المتاحة"

    def _postprocess_resources(
        self,
        items: list[dict[str, Any]],
        *,
        issue_title: str | None = None,
        source_collection: str,
        source_document: str,
        mbti_type: str | None = None,
        topic_title: str | None = None,
    ) -> list[dict[str, Any]]:
        processed: list[dict[str, Any]] = []
        for raw in items:
            raw_title = (raw.get("title") or "").strip()
            if not raw_title:
                continue

            url = raw.get("url") or self._extract_url(raw_title)
            title_without_url = self._remove_url_from_text(raw_title) if url else raw_title
            provisional_category = self._classify_resource(title_without_url or raw_title, url=url)
            clean_title = self._clean_resource_title(title_without_url or raw_title, provisional_category)

            if (not clean_title or clean_title == url) and processed:
                previous = processed[-1]
                if url and not previous.get("url"):
                    previous["url"] = url
                    previous["source_domain"] = self._extract_domain(url)
                    previous["category"] = self._classify_resource(previous.get("title", ""), url=url)
                    continue

            if not clean_title or self._looks_like_noise_resource(clean_title):
                continue

            category = self._classify_resource(clean_title, url=url)
            item = {
                "title": clean_title,
                "category": category,
                "url": url,
                "issue_title": raw.get("issue_title") or issue_title,
                "topic_title": raw.get("topic_title") or topic_title,
                "source_collection": source_collection,
                "source_domain": self._extract_domain(url),
                "source_document": source_document,
                "mbti_type": mbti_type,
            }

            if category == "general" and not url:
                continue
            if category in {"video", "podcast"} and not url:
                # keep in KB for tracing if useful later, but don't return by recommend unless URL required.
                pass
            processed.append(item)

        merged: list[dict[str, Any]] = []
        seen: set[str] = set()
        for item in processed:
            key = self._resource_key(item)
            if key in seen:
                continue
            seen.add(key)
            merged.append(item)
        return merged

    def _clean_resource_title(self, text: str, category: str) -> str:
        cleaned = re.sub(r"\s+", " ", text).strip("-•* ")
        cleaned = re.sub(r"^\(\d+\)\s*", "", cleaned)
        cleaned = re.sub(r"\s*-\s*YouTube$", "", cleaned, flags=re.I)
        cleaned = re.sub(r"\s*\|\s*YouTube$", "", cleaned, flags=re.I)

        quote_match = re.search(r"[\"“](.{4,120}?)[\"”]", cleaned)
        if not quote_match and "“" in cleaned:
            quote_match = re.search(r"“(.{4,120})", cleaned)
        if quote_match:
            quoted = quote_match.group(1).strip().strip("”\"")
            if 4 <= len(quoted) <= 120:
                cleaned = quoted

        for marker in self.DESCRIPTIVE_CUT_MARKERS:
            if marker in cleaned and len(cleaned.split()) > 10:
                cleaned = cleaned.split(marker, 1)[0].strip(" -—–:؛،")
                break
        if ": قناة" in cleaned:
            cleaned = cleaned.split(": قناة", 1)[0].strip()

        if ":" in cleaned:
            left, right = [part.strip() for part in cleaned.split(":", 1)]
            if any(left.startswith(prefix) for prefix in self.GENERIC_PREFIXES) and right:
                cleaned = right

        if category == "book" and any(channel in cleaned for channel in self.KNOWN_BOOK_CHANNELS):
            cleaned = re.split(r"\||-", cleaned)[0].strip()
        if category == "book":
            paren_match = re.search(r"\((كتاب[^)]{2,80})\)", cleaned)
            if paren_match:
                cleaned = paren_match.group(1).strip()

        if "|" in cleaned:
            parts = [part.strip() for part in cleaned.split("|") if part.strip()]
            if len(parts) >= 2 and len(" | ".join(parts[:2])) <= 110:
                cleaned = " | ".join(parts[:2])
            elif parts:
                cleaned = parts[0]

        if cleaned.startswith("مثل "):
            cleaned = re.sub(r"^مثل\s+", "", cleaned).strip()
        cleaned = re.sub(r"^سلسلة\s+", "", cleaned).strip()
        cleaned = re.sub(r"^حلقة\s*\((.*?)\)$", r"\1", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip("-—–:؛،. ")
        if category == "book" and len(cleaned) > 80:
            cleaned = re.split(r"[.!؟]", cleaned, maxsplit=1)[0].strip()
        if len(cleaned) > 110:
            sentence = re.split(r"[.!؟]", cleaned, maxsplit=1)[0].strip()
            cleaned = sentence if 8 <= len(sentence) <= 110 else cleaned[:107].rstrip() + "..."
        normalized_clean = self._normalize_label(cleaned)
        if normalized_clean in {self._normalize_label(x) for x in self.KNOWN_BOOK_CHANNELS + ["دوباميكافين", "بوكافيين"]}:
            return ""
        if cleaned.startswith("إذا ") or cleaned.startswith("ده ") or cleaned.startswith("لأن"):
            return ""
        return cleaned

    def _extract_domain(self, url: str | None) -> str | None:
        if not url:
            return None
        try:
            host = urlparse(url).netloc.lower().strip()
            return host or None
        except Exception:
            return None

    def _has_valid_url(self, url: str | None) -> bool:
        if not url or not isinstance(url, str):
            return False
        clean = url.strip().lower()
        return clean.startswith("http://") or clean.startswith("https://")

    def _dedupe_similar(self, items: list[str]) -> list[str]:
        result: list[str] = []
        for item in items:
            norm = normalize_text(item)
            if not norm:
                continue
            if any(lexical_overlap_score(norm, normalize_text(existing)) >= 0.72 for existing in result):
                continue
            result.append(item)
        return result

    def _slugify(self, text: str) -> str:
        normalized = self._normalize_label(text)
        return re.sub(r"\s+", "-", normalized)[:60]
